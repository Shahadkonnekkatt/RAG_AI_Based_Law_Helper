"""
doc_generator.py — LegalEase Legal Document Generator
Supports:
  1. Complaint Letter  — To SHO or Magistrate (BNSS 2023 s.173 / s.175)
  2. Legal Notice      — General civil grievance (BNS / Contract / Consumer)
  3. Cybercrime Complaint — To Cyber Cell (IT Act 2000 / BNS 2023)
"""

import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colour palette ─────────────────────────────────────────────────────────
NAVY = RGBColor(26,  42,  68)
GOLD = RGBColor(180, 130, 20)
GRAY = RGBColor(110, 110, 110)


# ── Shared helpers ──────────────────────────────────────────────────────────

def _validate_fields(data: dict, required: list):
    missing = [f for f in required if not str(data.get(f, "")).strip()]
    if missing:
        raise ValueError(f"Missing required fields: {', '.join(missing)}")


def _page_setup(doc, left=1.2, right=1.2, top=1.0, bottom=1.0):
    s = doc.sections[0]
    s.page_width    = Inches(8.27)
    s.page_height   = Inches(11.69)
    s.left_margin   = Inches(left)
    s.right_margin  = Inches(right)
    s.top_margin    = Inches(top)
    s.bottom_margin = Inches(bottom)


def _divider(doc, color="1A2A44", sz="6"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bt = OxmlElement("w:bottom")
    bt.set(qn("w:val"),   "single")
    bt.set(qn("w:sz"),    sz)
    bt.set(qn("w:space"), "1")
    bt.set(qn("w:color"), color)
    pBdr.append(bt)
    pPr.append(pBdr)


def _spacer(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)


def _para(doc, text="", bold=False, italic=False, size=11,
          color=None, center=False, align_right=False,
          left_indent=0.0, first_indent=0.0,
          space_before=4, space_after=4, justify=False):
    """Add a paragraph with a single run. Returns the paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align_right:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if left_indent:
        p.paragraph_format.left_indent = Inches(left_indent)
    if first_indent:
        p.paragraph_format.first_line_indent = Inches(first_indent)
    if text:
        r = p.add_run(str(text))
        r.bold   = bold
        r.italic = italic
        r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color
    return p


def _label_value(doc, label, value, size=11):
    """Bold label + plain value on one line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    lr = p.add_run(label + ": ")
    lr.bold = True
    lr.font.size = Pt(size)
    lr.font.color.rgb = NAVY
    vr = p.add_run(str(value).strip() if value else "—")
    vr.font.size = Pt(size)
    return p


def _section_heading(doc, text, size=11):
    """Bold section heading with thin rule below."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = NAVY
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bt = OxmlElement("w:bottom")
    bt.set(qn("w:val"),   "single")
    bt.set(qn("w:sz"),    "4")
    bt.set(qn("w:space"), "1")
    bt.set(qn("w:color"), "CCCCCC")
    pBdr.append(bt)
    pPr.append(pBdr)
    return p


def _save(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    data = buf.read()
    if len(data) < 100:
        raise RuntimeError("Document generation produced empty or corrupt output.")
    return data


# ── Document 1: Complaint Letter ───────────────────────────────────────────

def generate_complaint_letter(data: dict) -> bytes:
    """
    Formal complaint letter requesting FIR registration.
    Addressed to SHO (s.173 BNSS 2023) or Magistrate (s.175 BNSS 2023).
    """
    _validate_fields(data, [
        "complainant_name", "complainant_address",
        "incident_description", "place_of_incident",
    ])

    doc = Document()
    _page_setup(doc)

    now_str  = datetime.now().strftime("%d %B %Y")
    date_str = str(data.get("date", "")).strip() or now_str

    addressee_raw  = str(data.get("addressee_type", "SHO")).strip().upper()
    use_magistrate = any(x in addressee_raw for x in ("MAGISTRATE", "COURT", "JMFC"))

    # ── Date (right-aligned) ──────────────────────────────────────────────
    _para(doc, date_str, size=11, align_right=True, space_before=0, space_after=14)

    # ── Addressee block ───────────────────────────────────────────────────
    _para(doc, "To,", bold=True, size=11, space_before=0, space_after=2)

    if use_magistrate:
        court_name = str(data.get("court_name", "")).strip()
        district   = str(data.get("district",   "")).strip()
        addr_lines = [
            "The Judicial Magistrate (First Class),",
            (court_name + ",") if court_name else "____________________________,",
            district if district else "____________________________",
        ]
        authority_clause = (
            "respectfully submit this application under Section 175 of the "
            "Bharatiya Nagarik Suraksha Sanhita, 2023 (BNSS) requesting your "
            "Honour to direct the concerned police station to register a "
            "First Information Report in connection with the following matter."
        )
        subject_desc = "Application u/s 175 BNSS 2023 — Direction to Register FIR"
    else:
        ps_name  = str(data.get("police_station", "")).strip()
        district = str(data.get("district",       "")).strip()
        addr_lines = [
            "The Station House Officer,",
            ((ps_name + " Police Station,") if ps_name else "____________________ Police Station,"),
            district if district else "____________________________",
        ]
        authority_clause = (
            "hereby lodge this complaint and respectfully request you to register "
            "a First Information Report (FIR) under Section 173 of the Bharatiya "
            "Nagarik Suraksha Sanhita, 2023 (BNSS) in connection with the "
            "following matter."
        )
        subject_desc = "Complaint for Registration of FIR u/s 173 BNSS 2023"

    for line in addr_lines:
        _para(doc, line, size=11, left_indent=0.3, space_before=0, space_after=1)

    _spacer(doc)

    # ── Subject line ──────────────────────────────────────────────────────
    offence = str(data.get("offence_type", "")).strip()
    subject_full = subject_desc + (" regarding " + offence if offence else "")

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(4)
    p_sub.paragraph_format.space_after  = Pt(4)
    sr = p_sub.add_run("Subject: ")
    sr.bold = True
    sr.font.size = Pt(11)
    sr.font.color.rgb = NAVY
    sv = p_sub.add_run(subject_full)
    sv.font.size = Pt(11)
    sv.underline = True

    _divider(doc, "1A2A44", "4")
    _spacer(doc)

    # ── Salutation ────────────────────────────────────────────────────────
    _para(doc, "Respected Sir / Madam,", size=11, space_before=2, space_after=8)

    # ── Opening paragraph ─────────────────────────────────────────────────
    name   = str(data.get("complainant_name",    "")).strip()
    age    = str(data.get("complainant_age",     "")).strip()
    gender = str(data.get("complainant_gender",  "")).strip()
    addr   = str(data.get("complainant_address", "")).strip()
    phone  = str(data.get("complainant_phone",   "")).strip()

    desc_parts = []
    if age:    desc_parts.append(f"aged {age}")
    if gender: desc_parts.append(gender)
    desc_str  = (", " + ", ".join(desc_parts) + ",") if desc_parts else ","
    phone_str = f", contactable at {phone}" if phone else ""

    opening = (
        f"I, {name}{desc_str} residing at {addr}{phone_str}, "
        f"{authority_clause}"
    )
    _para(doc, opening, size=11, justify=True, first_indent=0.3,
          space_before=2, space_after=8)

    # ── Facts of the case ─────────────────────────────────────────────────
    _section_heading(doc, "FACTS OF THE CASE")

    inc_date = str(data.get("incident_date", "")).strip()
    inc_time = str(data.get("incident_time", "")).strip()
    place    = str(data.get("place_of_incident", "")).strip()

    if inc_date or place:
        parts = []
        if inc_date:
            dt = inc_date + (f" at approximately {inc_time}" if inc_time else "")
            parts.append(f"On {dt}")
        if place:
            parts.append(f"at {place}")
        intro = ", ".join(parts) + ", the following incident took place:"
        _para(doc, intro, size=11, left_indent=0.3, space_before=4, space_after=4)

    _para(doc, str(data.get("incident_description", "")).strip(),
          size=11, justify=True, left_indent=0.3, space_before=2, space_after=8)

    # ── Accused details ───────────────────────────────────────────────────
    acc_name = str(data.get("accused_name",    "")).strip()
    acc_addr = str(data.get("accused_address", "")).strip()

    _section_heading(doc, "DETAILS OF ACCUSED")
    if acc_name and acc_name.lower() not in ("unknown", "not known", ""):
        _label_value(doc, "Name",    acc_name)
        _label_value(doc, "Address", acc_addr if acc_addr else "Not known")
    else:
        _para(doc,
              "The identity of the accused is currently unknown and is requested "
              "to be ascertained during the course of investigation.",
              size=11, italic=True, left_indent=0.3,
              space_before=2, space_after=4)

    # ── Witnesses ─────────────────────────────────────────────────────────
    witnesses = str(data.get("witnesses", "")).strip()
    if witnesses and witnesses.lower() not in ("none", "no witnesses", "unknown", ""):
        _section_heading(doc, "WITNESSES")
        _para(doc, witnesses, size=11, left_indent=0.3, space_before=2, space_after=4)

    # ── Prayer ────────────────────────────────────────────────────────────
    _section_heading(doc, "PRAYER / RELIEF SOUGHT")

    relief_custom = str(data.get("relief_sought", "")).strip()
    if relief_custom:
        _para(doc, relief_custom, size=11, justify=True, left_indent=0.3,
              space_before=2, space_after=6)
    else:
        if use_magistrate:
            prayers = [
                "Direct the concerned police station to register an FIR under the appropriate sections of the Bharatiya Nyaya Sanhita, 2023.",
                "Direct a thorough and impartial investigation into the matter.",
                "Pass any other order as deemed fit in the interest of justice.",
            ]
        else:
            prayers = [
                "Register an FIR under the appropriate sections of the Bharatiya Nyaya Sanhita, 2023.",
                "Conduct a thorough investigation into the matter at the earliest.",
                "Take appropriate legal action against the accused.",
                "Provide me a copy of the FIR upon registration.",
            ]
        _para(doc, "I humbly request you to:", size=11, left_indent=0.3,
              space_before=4, space_after=2)
        for i, prayer in enumerate(prayers, 1):
            _para(doc, f"{i}. {prayer}", size=11, left_indent=0.5,
                  space_before=1, space_after=2)

    # ── Declaration ───────────────────────────────────────────────────────
    _spacer(doc)
    _para(doc,
          "I declare that the information furnished above is true and correct "
          "to the best of my knowledge and belief.",
          size=11, justify=True, first_indent=0.3, space_before=4, space_after=10)

    # ── Signature block ───────────────────────────────────────────────────
    _para(doc, "Yours faithfully,", size=11, space_before=2, space_after=1)
    _spacer(doc, 2)

    sig_p = doc.add_paragraph()
    sig_p.paragraph_format.space_before = Pt(1)
    sig_p.paragraph_format.space_after  = Pt(1)
    sig_r = sig_p.add_run(name)
    sig_r.bold = True
    sig_r.font.size = Pt(11)

    if addr:
        _para(doc, addr,             size=10, color=GRAY, space_before=1, space_after=1)
    if phone:
        _para(doc, phone,            size=10, color=GRAY, space_before=1, space_after=1)
    _para(doc, "Date: " + date_str,  size=10, color=GRAY, space_before=1, space_after=4)

    _divider(doc, "AAAAAA", "2")

    # ── Footer note ───────────────────────────────────────────────────────
    _para(doc,
          "Note: This is a citizen's request for FIR registration under Section 173, "
          "BNSS 2023. Please retain a signed copy of this letter. The FIR number will "
          "be assigned by the Station House Officer upon registration.",
          size=8, italic=True, color=GRAY, space_before=2, space_after=4)

    return _save(doc)


# ── Document 2: Legal Notice ───────────────────────────────────────────────

def generate_legal_notice(data: dict) -> bytes:
    """
    General Legal Notice covering money recovery, service deficiency,
    contract breach, property disputes, and other civil grievances.
    Advocate is optional — may be sent pro se.
    """
    _validate_fields(data, [
        "sender_name", "sender_address",
        "recipient_name", "recipient_address",
        "subject", "grievance_description", "relief_demanded",
    ])

    doc = Document()
    _page_setup(doc)

    now_str    = datetime.now().strftime("%d %B %Y")
    date_str   = str(data.get("date", "")).strip() or now_str
    reply_days = str(data.get("reply_days", "")).strip() or "15"

    advocate_name  = str(data.get("advocate_name",     "")).strip()
    advocate_addr  = str(data.get("advocate_address",  "")).strip()
    advocate_enrol = str(data.get("advocate_enrolment","")).strip()
    has_advocate   = bool(advocate_name)

    sender_name = str(data.get("sender_name",    "")).strip()
    sender_addr = str(data.get("sender_address", "")).strip()

    # ── Letterhead ────────────────────────────────────────────────────────
    if has_advocate:
        _para(doc, advocate_name.upper(), bold=True, size=15, center=True,
              color=NAVY, space_before=0, space_after=4)
        if advocate_enrol:
            _para(doc, "Bar Council Enrolment No.: " + advocate_enrol,
                  size=10, center=True, color=GRAY, space_before=0, space_after=3)
        if advocate_addr:
            _para(doc, advocate_addr, size=10, center=True,
                  color=GRAY, space_before=0, space_after=4)
        _divider(doc, "B89020", "6")
    else:
        _para(doc, sender_name.upper(), bold=True, size=13, center=True,
              color=NAVY, space_before=0, space_after=3)
        _para(doc, sender_addr, size=10, center=True,
              color=GRAY, space_before=0, space_after=4)
        _divider(doc, "1A2A44", "4")

    _spacer(doc)

    # ── Date (right) ──────────────────────────────────────────────────────
    _para(doc, "Date: " + date_str, size=11, align_right=True,
          space_before=2, space_after=12)

    # ── Addressee ─────────────────────────────────────────────────────────
    _para(doc, "To,", bold=True, size=11, space_before=0, space_after=2)
    _para(doc, str(data.get("recipient_name",    "")).strip(), size=11,
          left_indent=0.3, space_before=0, space_after=1)
    _para(doc, str(data.get("recipient_address", "")).strip(), size=11,
          left_indent=0.3, space_before=0, space_after=6)

    # ── Subject ───────────────────────────────────────────────────────────
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(4)
    p_sub.paragraph_format.space_after  = Pt(8)
    sr = p_sub.add_run("Subject: ")
    sr.bold = True; sr.font.size = Pt(11); sr.font.color.rgb = NAVY
    sv = p_sub.add_run(str(data.get("subject", "Legal Notice")).strip())
    sv.font.size = Pt(11); sv.underline = True

    _divider(doc)
    _spacer(doc)

    # ── Salutation ────────────────────────────────────────────────────────
    _para(doc, "Sir / Madam,", size=11, space_before=2, space_after=8)

    # ── Opening paragraph ─────────────────────────────────────────────────
    if has_advocate:
        opening = (
            f"I, the undersigned advocate, have been duly instructed and authorized "
            f"by my client, {sender_name}, residing at {sender_addr} "
            f"(hereinafter referred to as \"my client\"), to address you this Legal Notice."
        )
    else:
        opening = (
            f"I, {sender_name}, residing at {sender_addr}, "
            f"hereby serve upon you this Legal Notice drawing your attention to the "
            f"following grievance and calling upon you to address the same."
        )
    _para(doc, opening, size=11, justify=True, first_indent=0.3,
          space_before=2, space_after=8)

    # ── Grievance ─────────────────────────────────────────────────────────
    _section_heading(doc, "FACTS AND GRIEVANCE")
    _para(doc, str(data.get("grievance_description", "")).strip(),
          size=11, justify=True, left_indent=0.3, space_before=4, space_after=8)

    # ── Legal basis ───────────────────────────────────────────────────────
    legal_basis = str(data.get("legal_basis", "")).strip()
    if not legal_basis:
        client_ref = "my client" if has_advocate else "the undersigned"
        legal_basis = (
            f"Your aforesaid conduct and continued failure to address the above "
            f"grievance is in violation of the applicable provisions of law and has "
            f"caused substantial loss, hardship, and prejudice to {client_ref}."
        )
    _para(doc, legal_basis, size=11, justify=True, left_indent=0.3,
          space_before=2, space_after=8)

    # ── Demand ────────────────────────────────────────────────────────────
    _section_heading(doc, "DEMAND / RELIEF SOUGHT")
    _para(doc, str(data.get("relief_demanded", "")).strip(),
          size=11, justify=True, left_indent=0.3, space_before=4, space_after=8)

    # ── Warning ───────────────────────────────────────────────────────────
    client_or_i = "my client" if has_advocate else "I"
    warning = (
        f"You are hereby called upon to comply with the above demand within "
        f"{reply_days} days from the receipt of this notice, failing which "
        f"{client_or_i} shall be constrained to initiate appropriate legal "
        f"proceedings, both civil and criminal, before the competent forum or "
        f"Court of law, without any further notice, entirely at your cost and risk."
    )
    _para(doc, warning, size=11, justify=True, left_indent=0.3,
          space_before=4, space_after=8)

    _para(doc,
          "This notice is issued without prejudice to all other rights and "
          "remedies available under law.",
          size=10, italic=True, left_indent=0.3, space_before=2, space_after=12)

    # ── Signature ─────────────────────────────────────────────────────────
    _para(doc, "Yours faithfully,", size=11, space_before=2, space_after=2)
    _spacer(doc, 2)

    sig_p = doc.add_paragraph()
    sig_p.paragraph_format.space_before = Pt(2)
    sig_p.paragraph_format.space_after  = Pt(2)
    sig_name = advocate_name if has_advocate else sender_name
    sig_r = sig_p.add_run(sig_name)
    sig_r.bold = True
    sig_r.font.size = Pt(11)

    if has_advocate:
        _para(doc, f"Advocate for {sender_name}",
              size=10, color=GRAY, space_before=1, space_after=1)
    _para(doc, "Date: " + date_str,
          size=10, color=GRAY, space_before=1, space_after=4)

    return _save(doc)


# ── Document 3: Cybercrime Complaint ───────────────────────────────────────

def generate_cybercrime_complaint(data: dict) -> bytes:
    """
    Formal complaint to the Cyber Crime Cell / SP (Cyber).
    Filed under Information Technology Act, 2000 and BNS, 2023.
    """
    _validate_fields(data, [
        "complainant_name", "complainant_address",
        "complainant_phone", "incident_description",
    ])

    doc = Document()
    _page_setup(doc)

    now_str  = datetime.now().strftime("%d %B %Y")
    date_str = str(data.get("date", "")).strip() or now_str

    name        = str(data.get("complainant_name",     "")).strip()
    addr        = str(data.get("complainant_address",  "")).strip()
    phone       = str(data.get("complainant_phone",    "")).strip()
    email       = str(data.get("complainant_email",    "")).strip()
    age         = str(data.get("complainant_age",      "")).strip()
    gender      = str(data.get("complainant_gender",   "")).strip()
    crime_type  = str(data.get("crime_type",           "Cybercrime")).strip()
    platform    = str(data.get("platform",             "")).strip()
    inc_date    = str(data.get("incident_date",        "")).strip()
    amount_lost = str(data.get("amount_lost",          "")).strip()
    acc_details = str(data.get("accused_details",      "")).strip()
    evidence    = str(data.get("evidence_description", "")).strip()
    description = str(data.get("incident_description", "")).strip()
    district    = str(data.get("district",             "")).strip()
    state       = str(data.get("state",                "")).strip()

    # ── Document heading ──────────────────────────────────────────────────
    _para(doc, "COMPLAINT TO THE CYBER CRIME CELL",
          bold=True, size=14, center=True, color=NAVY,
          space_before=0, space_after=2)
    _para(doc,
          "Under the Information Technology Act, 2000 & Bharatiya Nyaya Sanhita, 2023",
          size=10, center=True, italic=True, color=GRAY,
          space_before=0, space_after=6)
    _divider(doc, "1A2A44", "6")
    _spacer(doc)

    # ── Date & Addressee ──────────────────────────────────────────────────
    _para(doc, "Date: " + date_str, size=11, align_right=True,
          space_before=0, space_after=10)
    _para(doc, "To,", bold=True, size=11, space_before=0, space_after=2)

    cy_addr_lines = [
        "The Station House Officer / Officer-in-Charge,",
        "Cyber Crime Police Station / Cyber Cell,",
    ]
    if district or state:
        cy_addr_lines.append(", ".join(filter(None, [district, state])))
    for line in cy_addr_lines:
        _para(doc, line, size=11, left_indent=0.3, space_before=0, space_after=1)

    _spacer(doc)

    # ── Subject ───────────────────────────────────────────────────────────
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(4)
    p_sub.paragraph_format.space_after  = Pt(8)
    sr = p_sub.add_run("Subject: ")
    sr.bold = True; sr.font.size = Pt(11); sr.font.color.rgb = NAVY
    sv = p_sub.add_run(
        f"Complaint regarding {crime_type} — Request for Investigation and Action")
    sv.font.size = Pt(11); sv.underline = True

    _divider(doc)
    _spacer(doc)

    # ── Salutation ────────────────────────────────────────────────────────
    _para(doc, "Respected Sir / Madam,", size=11, space_before=2, space_after=8)

    # ── Introduction ──────────────────────────────────────────────────────
    desc_parts = []
    if age:    desc_parts.append(f"aged {age}")
    if gender: desc_parts.append(gender)
    desc_str = (", " + ", ".join(desc_parts)) if desc_parts else ""
    contact  = phone + (f", {email}" if email else "")

    intro = (
        f"I, {name}{desc_str}, residing at {addr}, contactable at {contact}, "
        f"hereby file this complaint regarding a cybercrime committed against me "
        f"and request immediate investigation and appropriate legal action."
    )
    _para(doc, intro, size=11, justify=True, first_indent=0.3,
          space_before=2, space_after=8)

    # ── Section A: Nature of Complaint ────────────────────────────────────
    _section_heading(doc, "A.  NATURE OF COMPLAINT")
    _label_value(doc, "Type of Cybercrime", crime_type)
    _label_value(doc, "Platform / Medium",  platform if platform else "Not specified")
    _label_value(doc, "Date of Incident",   inc_date  if inc_date  else "Not specified")
    if amount_lost and amount_lost not in ("0", ""):
        _label_value(doc, "Amount Lost (₹)", amount_lost)

    # ── Section B: Description of Incident ───────────────────────────────
    _section_heading(doc, "B.  DESCRIPTION OF INCIDENT")
    _para(doc, description, size=11, justify=True, left_indent=0.3,
          space_before=4, space_after=8)

    # ── Section C: Accused Details ────────────────────────────────────────
    _section_heading(doc, "C.  DETAILS OF ACCUSED (IF KNOWN)")
    if acc_details and acc_details.lower() not in ("unknown", "not known", ""):
        _para(doc, acc_details, size=11, left_indent=0.3,
              space_before=4, space_after=4)
    else:
        _para(doc,
              "The identity of the accused is currently unknown. "
              "Investigation is requested to establish the same.",
              size=11, italic=True, left_indent=0.3,
              space_before=4, space_after=4)

    # ── Section D: Evidence ───────────────────────────────────────────────
    _section_heading(doc, "D.  EVIDENCE AVAILABLE")
    if evidence and evidence.lower() not in ("none", ""):
        _para(doc, evidence, size=11, left_indent=0.3,
              space_before=4, space_after=4)
    else:
        _para(doc,
              "Digital evidence including screenshots, transaction records, and "
              "communications are available and will be produced upon request.",
              size=11, italic=True, left_indent=0.3,
              space_before=4, space_after=4)

    # ── Section E: Legal Provisions ───────────────────────────────────────
    _section_heading(doc, "E.  APPLICABLE LEGAL PROVISIONS")
    for prov in [
        "Information Technology Act, 2000 — Sections 43, 66, 66B, 66C, 66D, 66E, 67, 67A as applicable",
        "Bharatiya Nyaya Sanhita, 2023 — Sections 318, 319 (Cheating / Fraud) as applicable",
    ]:
        _para(doc, "• " + prov, size=10, left_indent=0.4,
              space_before=2, space_after=2)

    # ── Section F: Prayer ─────────────────────────────────────────────────
    _section_heading(doc, "F.  PRAYER / RELIEF SOUGHT")
    prayers = [
        "Register an FIR and investigate this cybercrime matter on priority.",
        "Initiate appropriate legal action against the accused under the IT Act, 2000 and BNS, 2023.",
    ]
    if amount_lost and amount_lost not in ("0", ""):
        prayers.append(f"Assist in tracing and recovery of the financial loss of ₹{amount_lost}.")
    prayers.append("Issue me an acknowledgement and FIR copy upon registration.")

    _para(doc, "I humbly request you to:", size=11, left_indent=0.3,
          space_before=4, space_after=2)
    for i, p_text in enumerate(prayers, 1):
        _para(doc, f"{i}. {p_text}", size=11, left_indent=0.5,
              space_before=1, space_after=2)

    # ── Declaration ───────────────────────────────────────────────────────
    _spacer(doc)
    _para(doc,
          "I declare that the above information is true and correct to the "
          "best of my knowledge and belief.",
          size=11, justify=True, first_indent=0.3, space_before=4, space_after=10)

    # ── Signature ─────────────────────────────────────────────────────────
    _para(doc, "Yours sincerely,", size=11, space_before=2, space_after=2)
    _spacer(doc, 2)

    sig_p = doc.add_paragraph()
    sig_p.paragraph_format.space_before = Pt(2)
    sig_p.paragraph_format.space_after  = Pt(2)
    sig_r = sig_p.add_run(name)
    sig_r.bold = True
    sig_r.font.size = Pt(11)

    _para(doc, addr,             size=10, color=GRAY, space_before=1, space_after=1)
    _para(doc, contact,          size=10, color=GRAY, space_before=1, space_after=1)
    _para(doc, "Date: " + date_str, size=10, color=GRAY, space_before=1, space_after=8)

    _divider(doc, "AAAAAA", "2")

    _para(doc,
          "Important: You may also file this complaint online at www.cybercrime.gov.in "
          "or call the National Cyber Crime Helpline: 1930 (available 24×7).",
          size=8, italic=True, color=GRAY, space_before=6, space_after=4)

    return _save(doc)


# ── Document Registry ───────────────────────────────────────────────────────

DOCUMENT_TYPES = {
    "complaint_letter": {
        "label":       "Complaint Letter to SHO / Magistrate",
        "description": "Formal complaint requesting FIR registration under Section 173 / 175, BNSS 2023",
        "generator":   generate_complaint_letter,
        "fields": [
            {"name": "complainant_name",    "label": "Your Full Name",                       "required": True},
            {"name": "complainant_address", "label": "Your Full Address",                    "required": True},
            {"name": "incident_description","label": "Describe What Happened (in detail)",   "required": True,  "multiline": True},
            {"name": "place_of_incident",   "label": "Place Where Incident Occurred",        "required": True},
            {"name": "complainant_age",     "label": "Your Age",                             "required": False},
            {"name": "complainant_gender",  "label": "Your Gender",                          "required": False},
            {"name": "complainant_phone",   "label": "Your Phone Number",                    "required": False},
            {"name": "incident_date",       "label": "Date of Incident (DD/MM/YYYY)",        "required": False},
            {"name": "incident_time",       "label": "Approximate Time of Incident",         "required": False},
            {"name": "offence_type",        "label": "Nature / Type of Offence",             "required": False},
            {"name": "addressee_type",      "label": "Address To",                           "required": False,
             "type": "select", "options": ["SHO (Police Station)", "Magistrate (Court)"]},
            {"name": "police_station",      "label": "Police Station Name",                  "required": False},
            {"name": "court_name",          "label": "Court Name (if Magistrate)",           "required": False},
            {"name": "district",            "label": "District",                             "required": False},
            {"name": "accused_name",        "label": "Name of Accused (or Unknown)",         "required": False},
            {"name": "accused_address",     "label": "Address of Accused (or Unknown)",      "required": False},
            {"name": "witnesses",           "label": "Witnesses (Names & Details)",          "required": False, "multiline": True},
            {"name": "relief_sought",       "label": "Specific Relief Sought (optional)",    "required": False, "multiline": True},
            {"name": "date",                "label": "Date on Letter",                       "required": False},
        ],
    },

    "legal_notice": {
        "label":       "Legal Notice",
        "description": "Formal Legal Notice for civil grievances — money recovery, breach, dispute",
        "generator":   generate_legal_notice,
        "fields": [
            {"name": "sender_name",           "label": "Your Full Name (Notice Sender)",      "required": True},
            {"name": "sender_address",        "label": "Your Full Address",                   "required": True},
            {"name": "recipient_name",        "label": "Recipient Full Name",                 "required": True},
            {"name": "recipient_address",     "label": "Recipient Full Address",              "required": True},
            {"name": "subject",               "label": "Subject / Title of Notice",           "required": True},
            {"name": "grievance_description", "label": "Describe Your Grievance in Detail",   "required": True,  "multiline": True},
            {"name": "relief_demanded",       "label": "What You Demand from Recipient",      "required": True,  "multiline": True},
            {"name": "reply_days",            "label": "Days to Comply (default: 15)",        "required": False},
            {"name": "legal_basis",           "label": "Legal Basis / Law Violated (optional)","required": False, "multiline": True},
            {"name": "advocate_name",         "label": "Advocate Name (if applicable)",       "required": False},
            {"name": "advocate_address",      "label": "Advocate Office Address",             "required": False},
            {"name": "advocate_enrolment",    "label": "Bar Council Enrolment No.",           "required": False},
            {"name": "date",                  "label": "Date of Notice",                      "required": False},
        ],
    },

    "cybercrime_complaint": {
        "label":       "Cybercrime Complaint",
        "description": "Complaint to Cyber Cell under IT Act 2000 & BNS 2023",
        "generator":   generate_cybercrime_complaint,
        "fields": [
            {"name": "complainant_name",     "label": "Your Full Name",                        "required": True},
            {"name": "complainant_address",  "label": "Your Full Address",                     "required": True},
            {"name": "complainant_phone",    "label": "Your Phone Number",                     "required": True},
            {"name": "incident_description", "label": "Describe the Cybercrime in Detail",     "required": True, "multiline": True},
            {"name": "complainant_email",    "label": "Your Email Address",                    "required": False},
            {"name": "complainant_age",      "label": "Your Age",                              "required": False},
            {"name": "complainant_gender",   "label": "Your Gender",                           "required": False},
            {"name": "crime_type",           "label": "Type of Cybercrime",                    "required": False,
             "type": "select", "options": [
                 "Financial Fraud / Online Scam",
                 "Hacking / Unauthorized Access",
                 "Identity Theft",
                 "Online Harassment / Threats",
                 "Obscene Content / Privacy Violation",
                 "Phishing / Fake Website",
                 "Data Breach",
                 "Other",
             ]},
            {"name": "platform",             "label": "Platform / Website / App Involved",     "required": False},
            {"name": "incident_date",        "label": "Date of Incident (DD/MM/YYYY)",         "required": False},
            {"name": "amount_lost",          "label": "Amount Lost in ₹ (if financial fraud)", "required": False},
            {"name": "accused_details",      "label": "Accused Details (email/phone/username)","required": False, "multiline": True},
            {"name": "evidence_description", "label": "Evidence Available (screenshots, IDs)", "required": False, "multiline": True},
            {"name": "district",             "label": "District",                              "required": False},
            {"name": "state",                "label": "State",                                 "required": False},
            {"name": "date",                 "label": "Date on Complaint",                     "required": False},
        ],
    },
}


# ── Entry point ─────────────────────────────────────────────────────────────

def generate_document(doc_type: str, data: dict) -> tuple:
    """
    Main entry point.
    Returns (docx_bytes: bytes, filename: str).
    Raises ValueError for unknown doc_type.
    """
    if doc_type not in DOCUMENT_TYPES:
        raise ValueError(
            f"Unknown document type: '{doc_type}'. "
            f"Valid types: {list(DOCUMENT_TYPES.keys())}"
        )
    cfg        = DOCUMENT_TYPES[doc_type]
    docx_bytes = cfg["generator"](data)
    ts         = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename   = f"{doc_type}_{ts}.docx"
    return docx_bytes, filename